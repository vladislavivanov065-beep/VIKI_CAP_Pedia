from io import BytesIO

import pytest
from docx import Document

from apps.articles.document_import import parse_uploaded_document, split_document_into_blocks
from apps.attachments.exceptions import (
    AttachmentTooLargeError,
    InvalidAttachmentError,
    UnsupportedAttachmentFormatError,
)
from apps.attachments.tests.factories import make_pdf_bytes


def test_markdown_headings_split_txt_into_blocks():
    text = "# Введение\nПервый абзац.\n\n# Заключение\nВторой абзац."
    blocks = split_document_into_blocks(
        file_bytes=text.encode("utf-8"), extension="txt", fallback_title="Документ"
    )

    assert [b.title for b in blocks] == ["Введение", "Заключение"]
    assert blocks[0].content == "Первый абзац."
    assert blocks[1].content == "Второй абзац."


def test_content_before_first_heading_becomes_leading_block():
    text = "Вводный абзац без заголовка.\n\n# Раздел\nТекст раздела."
    blocks = split_document_into_blocks(
        file_bytes=text.encode("utf-8"), extension="txt", fallback_title="Документ"
    )

    assert blocks[0].title == "Документ"
    assert blocks[0].content == "Вводный абзац без заголовка."
    assert blocks[1].title == "Раздел"


def test_no_headings_produces_single_block_with_fallback_title():
    text = "Просто текст без каких-либо заголовков."
    blocks = split_document_into_blocks(
        file_bytes=text.encode("utf-8"), extension="txt", fallback_title="Документ"
    )

    assert len(blocks) == 1
    assert blocks[0].title == "Документ"
    assert blocks[0].content == text


def test_numbered_heading_heuristic_detected_in_plain_text():
    text = "1. Введение\nТекст первого раздела.\n2. Заключение\nТекст второго раздела."
    blocks = split_document_into_blocks(
        file_bytes=text.encode("utf-8"), extension="txt", fallback_title="Документ"
    )

    assert [b.title for b in blocks] == ["Введение", "Заключение"]


def test_all_caps_short_line_detected_as_heading():
    text = "ОБЩИЕ ПОЛОЖЕНИЯ\nТекст раздела про общие положения."
    blocks = split_document_into_blocks(
        file_bytes=text.encode("utf-8"), extension="txt", fallback_title="Документ"
    )

    assert blocks[0].title == "ОБЩИЕ ПОЛОЖЕНИЯ"
    assert blocks[0].content == "Текст раздела про общие положения."


def test_sentence_ending_punctuation_prevents_heading_misdetection():
    text = "1. Это обычное предложение, а не заголовок.\nВторая строка абзаца."
    blocks = split_document_into_blocks(
        file_bytes=text.encode("utf-8"), extension="txt", fallback_title="Документ"
    )

    assert len(blocks) == 1
    assert "Это обычное предложение" in blocks[0].content


def test_docx_heading_style_detected():
    document = Document()
    document.add_heading("Первый раздел", level=1)
    document.add_paragraph("Текст первого раздела.")
    document.add_heading("Второй раздел", level=2)
    document.add_paragraph("Текст второго раздела.")
    buffer = BytesIO()
    document.save(buffer)

    blocks = split_document_into_blocks(
        file_bytes=buffer.getvalue(), extension="docx", fallback_title="Документ"
    )

    assert [b.title for b in blocks] == ["Первый раздел", "Второй раздел"]
    assert blocks[0].content == "Текст первого раздела."
    assert blocks[1].content == "Текст второго раздела."


def test_pdf_with_no_extractable_text_produces_single_empty_block():
    blocks = split_document_into_blocks(
        file_bytes=make_pdf_bytes(), extension="pdf", fallback_title="Документ"
    )

    assert len(blocks) == 1
    assert blocks[0].title == "Документ"


def test_parse_uploaded_document_rejects_unsupported_extension():
    with pytest.raises(UnsupportedAttachmentFormatError):
        parse_uploaded_document(file_obj=BytesIO(b"data"), original_filename="report.exe")


def test_parse_uploaded_document_rejects_empty_file():
    with pytest.raises(InvalidAttachmentError):
        parse_uploaded_document(file_obj=BytesIO(b""), original_filename="report.txt")


def test_parse_uploaded_document_rejects_oversized_file(settings):
    settings.MAX_ATTACHMENT_SIZE_MB = 0
    with pytest.raises(AttachmentTooLargeError):
        parse_uploaded_document(file_obj=BytesIO(b"x" * 1024), original_filename="report.txt")


def test_parse_uploaded_document_rejects_corrupt_docx():
    with pytest.raises(InvalidAttachmentError):
        parse_uploaded_document(
            file_obj=BytesIO(b"not a real docx file"), original_filename="report.docx"
        )


def test_parse_uploaded_document_derives_fallback_title_from_filename():
    blocks = parse_uploaded_document(
        file_obj=BytesIO("Просто текст.".encode("utf-8")),
        original_filename="Регламент отпусков.txt",
    )

    assert blocks[0].title == "Регламент отпусков"
